from flask import Flask, request, render_template, send_file
import os
import easyocr
import pytesseract
from PIL import Image
from werkzeug.utils import secure_filename
import io

# For DOCX export
from docx import Document

# For PDF export
from fpdf import FPDF

# For Word/Excel/PDF conversion to image/text, need extra libraries (see notes below)

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'pdf', 'docx', 'xlsx'}

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_to_text(file_path, method):
    ext = file_path.rsplit('.', 1)[-1].lower()
    text = ""
    # Handle images directly
    if ext in ['png', 'jpg', 'jpeg']:
        if method == 'easyocr':
            reader = easyocr.Reader(['en'])
            results = reader.readtext(file_path, detail=0)
            text = "\n".join(results)
        elif method == 'tesseract':
            text = pytesseract.image_to_string(Image.open(file_path))
    # Handle PDF, DOCX, XLSX (requires additional libraries)
    elif ext == 'pdf':
        # Convert PDF pages to images
        from pdf2image import convert_from_path
        pages = convert_from_path(file_path)
        for page in pages:
            if method == 'easyocr':
                reader = easyocr.Reader(['en'])
                results = reader.readtext(page, detail=0)
                text += "\n".join(results) + "\n"
            elif method == 'tesseract':
                text += pytesseract.image_to_string(page) + "\n"
    elif ext == 'docx':
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif ext == 'xlsx':
        import pandas as pd
        df = pd.read_excel(file_path)
        text = df.to_string(index=False)
    return text

def save_text(text, save_format):
    if save_format == 'txt':
        return io.BytesIO(text.encode('utf-8')), 'text/plain'
    elif save_format == 'docx':
        doc = Document()
        doc.add_paragraph(text)
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif save_format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in text.split('\n'):
            pdf.cell(0, 10, line, ln=True)
        f = io.BytesIO()
        pdf.output(f)
        f.seek(0)
        return f, 'application/pdf'
    else:
        return io.BytesIO(text.encode('utf-8')), 'text/plain'

@app.route("/", methods=["GET", "POST"])
def upload_file():
    easyocr_text = tesseract_text = filename = None
    save_format = request.form.get("save_format") or "txt"
    ocr_method = request.form.get("ocr_method") or "easyocr"
    download_url = None

    if request.method == "POST":
        file = request.files.get("file")
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(filepath)

            # Run both OCRs if file is image/pdf
            easyocr_text = convert_to_text(filepath, "easyocr")
            tesseract_text = convert_to_text(filepath, "tesseract")

            # Save last OCR result in session for download
            request.environ['easyocr_text'] = easyocr_text
            request.environ['tesseract_text'] = tesseract_text
            request.environ['filepath'] = filepath

            selected_text = easyocr_text if ocr_method == "easyocr" else tesseract_text
            file_obj, mime_type = save_text(selected_text, save_format)
            file_obj.seek(0)
            # Save to a temp file for download link
            download_path = os.path.join(app.config["UPLOAD_FOLDER"], f"result.{save_format}")
            with open(download_path, "wb") as f:
                f.write(file_obj.read())
            download_url = url_for('download_file', filename=f"result.{save_format}")

    return render_template("upload.html",
                           filename=filename,
                           easyocr_text=easyocr_text,
                           tesseract_text=tesseract_text,
                           ocr_method=ocr_method,
                           save_format=save_format,
                           download_url=download_url)

@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    # Detect mimetype by file extension
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext == "pdf":
        mimetype = "application/pdf"
    elif ext == "docx":
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        mimetype = "text/plain"
    return send_file(file_path, as_attachment=True, mimetype=mimetype)

if __name__ == "__main__":
    app.run(debug=True)
